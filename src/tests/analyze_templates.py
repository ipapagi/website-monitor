from docx import Document
from openpyxl import load_workbook

# Read DOCX
docx_path = r'data\templates\ΗΜΕΡΟΜΗΝΙΑ_Αναφορά Εισερχομένων Αιτήσεων - για AI.docx'
doc = Document(docx_path)
print('=== DOCX STRUCTURE ===')
print(f'Total Paragraphs: {len(doc.paragraphs)}')
print(f'Total Tables: {len(doc.tables)}\n')

# Show all text
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f'P{i}: {para.text}')

# Show tables
for t_idx, table in enumerate(doc.tables):
    print(f'\n--- TABLE {t_idx} ---')
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f'Row {row_idx}: {cells}')

# Read XLSX
print('\n\n=== XLSX STRUCTURE ===')
xlsx_path = r'data\templates\ΓΔ_Γενική Διεύθυνση Αγροτικής Οικονομίας και Κτηνιατρικής_03-02-2026 ΟτΠ_ΓΔΑΟΚ 1 - Δείγμα ΣΗΔΕ για ΑΙ.xlsx'
wb = load_workbook(xlsx_path)
print(f'Sheet names: {wb.sheetnames}\n')

for sheet_name in wb.sheetnames:
    ws = wb[sheet_name]
    print(f'\nSheet: {sheet_name}')
    print(f'Dimensions: {ws.dimensions}')
    print('First 15 rows:')
    for i, row in enumerate(list(ws.iter_rows(values_only=True))[:15]):
        print(f'  Row {i+1}: {row}')
