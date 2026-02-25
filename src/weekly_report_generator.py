"""
Δημιουργία εβδομαδιαίας αναφοράς για ΣΗΔΕ αιτήσεις
Χρήση: python weekly_report_generator.py --date=2026-02-02
Για την ημερομηνία που δίνεις, παίρνει την προηγούμενη εβδομάδα (Δευτέρα-Κυριακή)
"""
import os
import sys
import argparse
import re
from datetime import datetime, timedelta

from docx import Document
from openpyxl import load_workbook

from config import get_project_root, INCOMING_DEFAULT_PARAMS
from utils import load_config
from monitor import PKMMonitor
from incoming import fetch_incoming_records, simplify_incoming_records
from api import enrich_record_details


def get_week_boundaries(date_str):
    """
    Παίρνει μια ημερομηνία και επιστρέφει την προηγούμενη εβδομάδα (Δευτέρα-Κυριακή)
    Π.χ., αν δώσεις 2026-02-02, θα πάρεις 2026-01-26 (Δευτέρα) ως 2026-02-01 (Κυριακή)
    """
    given_date = datetime.strptime(date_str, "%Y-%m-%d")
    
    # Βρες την Δευτέρα της παρούσας εβδομάδας
    days_since_monday = given_date.weekday()  # Monday=0, Sunday=6
    monday_of_this_week = given_date - timedelta(days=days_since_monday)
    
    # Η προηγούμενη εβδομάδα: Δευτέρα ως Κυριακή
    sunday_of_prev_week = monday_of_this_week - timedelta(days=1)
    monday_of_prev_week = sunday_of_prev_week - timedelta(days=6)
    
    return (
        monday_of_prev_week.strftime("%Y-%m-%d"),
        sunday_of_prev_week.strftime("%Y-%m-%d")
    )


def _ensure_logged_in(monitor: PKMMonitor):
    """Κάνει login και φορτώνει αρχική σελίδα αν χρειάζεται"""
    if not monitor.logged_in and not monitor.login():
        raise RuntimeError("Αποτυχία login")
    if not monitor.main_page_loaded:
        monitor.load_main_page()


def fetch_incoming_from_api(monitor: PKMMonitor, config: dict):
    """Ανακτά και απλοποιεί τις εγγραφές εισερχομένων από το API.
    
    Εξαιρεί:
    - Καταγγελίες (document_category = 'Καταγγελία')
    - Συμπληρωματικά Αιτήματος (document_category = 'Συμπληρωματικά Αιτήματος')
    
    ΠΡΟΣΟΧΗ: Οι εξαιρετικές εγγραφές δεν περιλαμβάνονται στις αναφορές, αλλά
    περιέχονται στο API snapshot για ιστορικό.
    """
    incoming_params = config.get('incoming_api_params', INCOMING_DEFAULT_PARAMS).copy()
    data = fetch_incoming_records(monitor, incoming_params)
    if not data or not data.get('success'):
        return []
    records = simplify_incoming_records(data.get('data', []))
    
    # Φιλτράρισμα: Αποκλεισμός καταγγελιών και Συμπληρωματικών Αιτημάτων
    excluded = {'καταγγελία', 'συμπληρωματικά αιτήματος'}
    filtered_records = [
        rec for rec in records 
        if rec.get('document_category', '').strip().lower() not in excluded
    ]
    
    excluded_count = len(records) - len(filtered_records)
    excluded_categories = {}
    for rec in records:
        if rec.get('document_category', '').strip().lower() in excluded:
            cat = rec.get('document_category', '').strip()
            excluded_categories[cat] = excluded_categories.get(cat, 0) + 1
    
    if excluded_categories:
        print(f"\n📊 Εγγραφές που εξαιρέθηκαν από τις αναφορές:")
        for category, count in sorted(excluded_categories.items()):
            print(f"   🚫 {count} εγγραφές: {category}")
        print(f"   (Αυτές δεν θα εμφανιστούν στα DOCX/XLSX, αλλά υπάρχουν στα snapshots για ιστορικό)")
    
    # Το enrichment θα γίνει στο main() για όλα τα records
    return filtered_records


def filter_records_by_week_end(records, sunday_str):
    """
    Φιλτράρει records για να εξαιρέσει αιτήσεις που υποβλήθηκαν ΜΕΤΑ το τέλος της εβδομάδας.
    Αυτό είναι σημαντικό όταν παράγουμε αναφορά για παλιές εβδομάδες.
    """
    sunday = datetime.strptime(sunday_str, "%Y-%m-%d")
    # Κάνε το inclusive - μέχρι τέλος της Κυριακής (23:59:59)
    week_end_cutoff = sunday + timedelta(days=1)  # Αρχή της επόμενης Δευτέρας
    
    filtered = []
    excluded_count = 0
    
    for rec in records:
        submitted_at_str = rec.get('submitted_at', '').strip()
        if not submitted_at_str:
            # Χωρίς ημερομηνία -> το κρατάμε
            filtered.append(rec)
            continue
        
        try:
            # Προσπάθησε διάφορες μορφές ημερομηνιών
            if 'T' in submitted_at_str:
                submitted_date = datetime.fromisoformat(submitted_at_str.replace('Z', '+00:00'))
            elif ' ' in submitted_at_str and len(submitted_at_str) > 10:
                submitted_date = datetime.strptime(submitted_at_str.split('.')[0], "%Y-%m-%d %H:%M:%S")
            elif len(submitted_at_str) == 10:
                submitted_date = datetime.strptime(submitted_at_str, "%Y-%m-%d")
            elif '/' in submitted_at_str:
                submitted_date = datetime.strptime(submitted_at_str, "%d/%m/%Y")
            else:
                filtered.append(rec)
                continue
            
            # Κράτα μόνο αιτήσεις που υποβλήθηκαν ΕΩΣ το τέλος της εβδομάδας
            if submitted_date < week_end_cutoff:
                filtered.append(rec)
            else:
                excluded_count += 1
        except (ValueError, AttributeError):
            # Αν δεν μπορούμε να parse την ημερομηνία, κρατάμε το record
            filtered.append(rec)
    
    if excluded_count > 0:
        print(f"\n⚠️  Εξαιρέθηκαν {excluded_count} νεότερες αιτήσεις (υποβλήθηκαν μετά το τέλος της εβδομάδας)")
    
    return filtered


def filter_out_settled_from_records(records, monday_str, general_directorate, report_date_str):
    """Αφαιρεί διεκπεραιωμένες πού έγιναν ΠΡΙΝ την ημερομηνία αναφοράς."""
    try:
        from settled_cases import fetch_settled_cases
        
        project_root = get_project_root()
        config = load_config(os.path.join(project_root, 'config', 'config.yaml'))
        monitor = PKMMonitor(
            base_url=config.get('base_url', 'https://shde.pkm.gov.gr'),
            urls=config.get('urls', {}), api_params=config.get('api_params', {}),
            login_params=config.get('login_params', {}), check_interval=config.get('check_interval', 300),
            username=config.get('username'), password=config.get('password'),
            session_cookies=config.get('session_cookies'))
        
        if not monitor.logged_in and not monitor.login():
            return records
        
        settled_result = fetch_settled_cases(monitor)
        settled_records = settled_result.get('data', []) if settled_result else []
        if not settled_records:
            return records
        
        report_date = datetime.strptime(report_date_str, "%Y-%m-%d")
        active_records = []
        removed_count = 0
        marked_count = 0
        removed_cases = []  # Track which cases were removed
        
        # Calculate sunday from monday
        monday_date = datetime.strptime(monday_str, "%Y-%m-%d")
        sunday_date = monday_date + timedelta(days=6)
        sunday_str = sunday_date.strftime("%Y-%m-%d")
        
        # Build lookup dict using protocol_number format (same as xls_export.py)
        # Maps "YYYY/CASE_ID" -> {'settled_date': ..., 'assigned_employee': ...}
        settled_by_protocol = {}
        for sr in settled_records:
            protocol_num = str(sr.get('W001_P_FLD2', '')).strip()
            settled_date_str = str(sr.get('W001_P_FLD3', '')).strip()
            assigned_employee = str(sr.get('W001_P_FLD10', '')).strip()
            
            if protocol_num:
                settled_by_protocol[protocol_num] = {
                    'settled_date': settled_date_str,
                    'assigned_employee': assigned_employee
                }
        
        print(f"   [DEBUG] Loaded {len(settled_by_protocol)} settled cases")
        print(f"   [DEBUG] Checking {len(records)} incoming records from {monday_str} to {sunday_str}")
        print(f"   [DEBUG] Report date: {report_date_str}")
        
        match_count = 0
        
        for rec in records:
            # Try matching using YEAR/CASE_ID format (same as xls_export.py)
            submission_year = str(rec.get('submission_year', '')).strip()
            case_id = str(rec.get('case_id', '')).strip()
            
            # Primary: Try YEAR/CASE_ID lookup
            settled = None
            lookup_key = None
            if submission_year and case_id:
                lookup_key = f"{submission_year}/{case_id}"
                settled = settled_by_protocol.get(lookup_key)
                if settled:
                    match_count += 1
            
            # Fallback: Try direct protocol_number from record
            if not settled:
                protocol_num = str(rec.get('protocol_number', '')).strip()
                if protocol_num:
                    settled = settled_by_protocol.get(protocol_num)
                    if settled:
                        match_count += 1
            
            # Fallback: Try extracting from related_case (old method)
            if not settled:
                related_case = rec.get('related_case', '')
                if related_case:
                    import re
                    match = re.search(r'(\d{4}/\d+)', str(related_case))
                    if match:
                        case_code = match.group(1).strip().upper()
                        settled = settled_by_protocol.get(case_code)
                        if settled:
                            match_count += 1
            
            if not settled:
                active_records.append(rec)
                continue
            
            completion_date_str = settled.get('settled_date', '')
            if not completion_date_str:
                active_records.append(rec)
                continue
            
            # Parse completion date - support multiple formats
            completion_date = None
            try:
                # Try ISO format first (YYYY-MM-DD or with time)
                if 'T' in completion_date_str:
                    completion_date = datetime.fromisoformat(completion_date_str.replace('Z', '+00:00')).date()
                elif '-' in completion_date_str:
                    date_part = completion_date_str.split(' ')[0].split('T')[0]
                    # Check if DD-MM-YYYY or YYYY-MM-DD
                    parts = date_part.split('-')
                    if len(parts) == 3:
                        if len(parts[0]) == 4:  # YYYY-MM-DD
                            completion_date = datetime.strptime(date_part, "%Y-%m-%d").date()
                        else:  # DD-MM-YYYY
                            completion_date = datetime.strptime(date_part, "%d-%m-%Y").date()
                elif '/' in completion_date_str:
                    # Try DD/MM/YYYY
                    date_part = completion_date_str.split(' ')[0]
                    completion_date = datetime.strptime(date_part, "%d/%m/%Y").date()
            except (ValueError, AttributeError) as e:
                print(f"   [ERROR] Case {case_id}: Failed to parse settled_date '{completion_date_str}' - {e}")
                active_records.append(rec)
                continue
            
            if not completion_date:
                active_records.append(rec)
                continue
            
            if completion_date < report_date.date():
                # Settled before report date -> REMOVE
                print(f"   [REMOVED] Case {case_id}: Settled on {completion_date} (before report date {report_date.date()})")
                removed_count += 1
                removed_cases.append(case_id)
                continue
            else:
                # Settled after or on report date - check if submitted this week
                submitted_at_str = rec.get('submitted_at', '')
                if submitted_at_str:
                    try:
                        if 'T' in submitted_at_str:
                            submitted_date = datetime.fromisoformat(submitted_at_str.replace('Z', '+00:00')).date()
                        else:
                            date_part = submitted_at_str.split(' ')[0]
                            submitted_date = datetime.strptime(date_part, "%Y-%m-%d").date()
                        
                        if monday_date.date() <= submitted_date <= sunday_date.date():
                            # New this week and settled -> MARK
                            rec['_settled_status'] = 'Διεκπεραιωμένη'
                            marked_count += 1
                    except (ValueError, AttributeError) as e:
                        pass
                
                # Keep it
                active_records.append(rec)
        
        if removed_count > 0:
            print(f"   Αφαιρέθησαν {removed_count} διεκπεραιωμένες (ΠΡΙΝ την {report_date_str})")
            if removed_cases:
                # Show first 10 case IDs
                sample = removed_cases[:10]
                more = f" και {len(removed_cases)-10} ακόμα" if len(removed_cases) > 10 else ""
                print(f"      Υποθέσεις: {', '.join(sample)}{more}")
        if marked_count > 0:
            print(f"   → {marked_count} νέες διεκπεραιωμένες σε αυτή την εβδομάδα")
            # Λίστα με τις συγκεκριμένες υποθέσεις
            settled_this_week_list = [r.get('case_id', 'N/A') for r in active_records if r.get('_settled_status') == 'Διεκπεραιωμένη']
            if settled_this_week_list:
                print(f"      Υποθέσεις: {', '.join(settled_this_week_list)}")
        print(f"   [DEBUG SUMMARY] Matched: {match_count}, Removed (before {report_date_str}): {removed_count}, Settled this week: {marked_count}")
        
        return active_records
    except Exception as e:
        print(f"   ΣΦΑΛΜΑ στο φιλτράρισμα: {e}")
        import traceback
        traceback.print_exc()
        return records


def categorize_records(records, monday_str, sunday_str):
    """
    Κατηγοριοποιεί τις αιτήσεις σε:
    - new_this_week: Νέες της εβδομάδας (submitted_at μεταξύ Δευτέρας-Κυριακής)
    - old: Παλιότερες (submitted_at πριν από τη Δευτέρα)
    - new_assigned: Νέες με ανάθεση
    - new_unassigned: Νέες χωρίς ανάθεση
    
    ΣΗΜΕΙΩΣΗ: Αυτή η συνάρτηση θεωρεί ότι τα records έχουν ήδη φιλτραριστεί
    να εξαιρούν νεότερες αιτήσεις (μέσω filter_records_by_week_end)
    """
    monday = datetime.strptime(monday_str, "%Y-%m-%d")
    sunday = datetime.strptime(sunday_str, "%Y-%m-%d") + timedelta(days=1)  # Κάνε το inclusive
    
    all_records = []
    seen_case_ids = set()
    for rec in records:
        case_id = rec.get('case_id')
        if case_id and case_id in seen_case_ids:
            continue
        if case_id:
            seen_case_ids.add(case_id)
        all_records.append(rec)
    
    new_this_week = []
    old = []
    
    for rec in all_records:
        submitted_at_str = rec.get('submitted_at', '').strip()
        if not submitted_at_str:
            # Αν δεν έχει submitted_at, χρησιμοποίησε το σήμερα (ή αγνόησε)
            old.append(rec)
            continue
        
        try:
            # Προσπάθησε διάφορες μορφές ημερομηνιών
            if 'T' in submitted_at_str:
                # ISO format με T: 2026-02-03T18:13:35.727528
                submitted_date = datetime.fromisoformat(submitted_at_str.replace('Z', '+00:00')).date()
            elif ' ' in submitted_at_str and len(submitted_at_str) > 10:
                # Format με κενό: 2026-02-03 18:13:35.727528
                date_part = submitted_at_str.split(' ')[0]
                submitted_date = datetime.strptime(date_part, "%Y-%m-%d").date()
            elif len(submitted_at_str) == 10:  # YYYY-MM-DD
                submitted_date = datetime.strptime(submitted_at_str, "%Y-%m-%d").date()
            elif '/' in submitted_at_str:  # DD/MM/YYYY ή MM/DD/YYYY
                submitted_date = datetime.strptime(submitted_at_str, "%d/%m/%Y").date()
            else:
                old.append(rec)
                continue
            
            if monday.date() <= submitted_date <= sunday.date():
                new_this_week.append(rec)
            else:
                old.append(rec)
        except (ValueError, AttributeError):
            old.append(rec)
    
    # Αφαίρεση διεκπεραιωμένων αιτήσεων από νέες (αιτήσεις που υποβλήθηκαν και διεκπεραιώθηκαν την ίδια εβδομάδα)
    new_this_week_active = [r for r in new_this_week if r.get('_settled_status') != 'Διεκπεραιωμένη']
    settled_this_week_count = len(new_this_week) - len(new_this_week_active)
    
    if settled_this_week_count > 0:
        print(f"   → Αφαιρέθησαν {settled_this_week_count} νέες αιτήσεις που ήδη διεκπεραιώθηκαν (δεν θα εμφανιστούν στα ΝΕΑ)")
    
    # Διαχωρισμός νέων αιτήσεων σε "με ανάθεση" (charged) και "χωρίς ανάθεση" (uncharged)
    new_assigned = [r for r in new_this_week_active if r.get('_charge', {}).get('charged')]
    new_unassigned = [r for r in new_this_week_active if not r.get('_charge', {}).get('charged')]
    
    return {
        'new_assigned': new_assigned,
        'new_unassigned': new_unassigned,
        'old': old,
        'all': all_records
    }


def group_by_general_directorate(records):
    """Ομαδοποιεί τις εγγραφές ανά Γενική Δ/νση."""
    groups = {}
    for rec in records:
        key = (rec.get('general_directorate') or '').strip()
        if not key:
            key = 'ΧΩΡΙΣ ΓΕΝΙΚΗ ΔΙΕΥΘΥΝΣΗ'
        groups.setdefault(key, []).append(rec)
    return groups


def print_statistics_table(groups, monday_str, sunday_str):
    """Εμφανίζει πίνακα στατιστικών με νέες και παλιές αιτήσεις ανά Γενική Δ/νση."""
    print("\n" + "="*90)
    print("📊 ΣΤΑΤΙΣΤΙΚΑ ΑΙΤΗΣΕΩΝ ΑΝΑ ΓΕΝΙΚΗ ΔΙΕΥΘΥΝΣΗ")
    print("="*90)
    
    # Συλλογή στατιστικών
    stats = []
    total_new = 0
    total_old = 0
    
    for general_directorate, recs in sorted(groups.items()):
        categorized = categorize_records(recs, monday_str, sunday_str)
        new_count = len(categorized['new_assigned']) + len(categorized['new_unassigned'])
        old_count = len(categorized['old'])
        
        # Συντομογραφία του ονόματος για καλύτερη εμφάνιση
        short_name = general_directorate.replace('ΓΕΝΙΚΗ ΔΙΕΥΘΥΝΣΗ ', '').replace('ΚΑΙ ', '& ')
        
        stats.append((short_name, new_count, old_count))
        total_new += new_count
        total_old += old_count
    
    # Κεφαλίδα πίνακα
    print(f"{'Γενική Διεύθυνση':<50} {'ΝΕΑ (εβδομάδας)':>15} {'ΠΑΛΙΑ (προηγούμενες)':>20}")
    print("-" * 90)
    
    # Γραμμές δεδομένων
    for short_name, new_count, old_count in stats:
        print(f"{short_name:<50} {new_count:>15} {old_count:>20}")
    
    # Σύνολα
    print("-" * 90)
    print(f"{'ΣΥΝΟΛΟ':<50} {total_new:>15} {total_old:>20}")
    print("="*90 + "\n")


def sanitize_filename_component(value):
    """Καθαρίζει κείμενο για χρήση σε όνομα αρχείου (Windows-safe).
    
    Keeps Greek characters but removes invalid filename characters.
    """
    cleaned = (value or '').strip()
    # Remove only invalid Windows filename characters
    cleaned = re.sub(r'[<>:"/\\|?*]', '', cleaned)
    # Replace spaces with underscores
    cleaned = re.sub(r'\s+', '_', cleaned)
    # Remove multiple consecutive underscores
    cleaned = re.sub(r'_+', '_', cleaned)
    return cleaned or 'ΧΩΡΙΣ_ΓΕΝΙΚΗ_ΔΙΕΥΘΥΝΣΗ'


def create_docx_report(template_path, output_path, report_date_str, categorized_data, general_directorate):
    """
    Δημιουργεί DOCX από το πρότυπο με τα σύνολα
    """
    doc = Document(template_path)
    
    # Υπολογισμός σύνολων
    new_assigned_count = len(categorized_data['new_assigned'])
    new_unassigned_count = len(categorized_data['new_unassigned'])
    old_count = len(categorized_data['old'])
    total_count = new_assigned_count + new_unassigned_count + old_count
    
    # Εξαγωγή ημερομηνιών εβδομάδας
    monday, sunday = get_week_boundaries(report_date_str)
    
    # Μετατροπή σε ημερομηνίες για εμφάνιση
    monday_date = datetime.strptime(monday, "%Y-%m-%d")
    sunday_date = datetime.strptime(sunday, "%Y-%m-%d")
    
    monday_gr = monday_date.strftime("%d-%m-%Y")
    sunday_gr = sunday_date.strftime("%d-%m-%Y")
    report_date_gr = datetime.strptime(report_date_str, "%Y-%m-%d").strftime("%d-%m-%Y")
    
    # Αντικατάστασης κειμένου στις παραγράφους
    for paragraph in doc.paragraphs:
        text = paragraph.text
        
        # Αντικατάστασης για τις ημερομηνίες της εβδομάδας
        if 'ΔΕΥΤΕΡΑ' in text and 'ΚΥΡΙΑΚΗ' in text:
            text = text.replace('ΔΕΥΤΕΡΑ', monday_gr).replace('ΚΥΡΙΑΚΗ', sunday_gr)
        if 'ΠΡΟΗΓΟΥΜΕΝΗ' in text:
            text = text.replace('ΠΡΟΗΓΟΥΜΕΝΗ', sunday_gr)
        
        # Αντικατάστασης για τα σύνολα
        if 'ΠΑΛΙΑ' in text:
            text = text.replace('ΠΑΛΙΑ', str(old_count))
        if 'Χωρίς Ανάθεση' in text:
            text = text.replace('Χωρίς Ανάθεση', f'είναι χωρίς ανάθεση, σύνολο: {new_unassigned_count}')
        if 'Ανάθεση' in text and 'Χωρίς' not in text:
            text = text.replace('Ανάθεση', f'είναι με ανάθεση, σύνολο: {new_assigned_count}')
        if 'ΕΔΩ' in text:
            text = text.replace('ΕΔΩ', str(total_count))
        if '[ΓΕΝΙΚΗ_ΔΙΕΥΘΥΝΣΗ]' in text:
            text = text.replace('[ΓΕΝΙΚΗ_ΔΙΕΥΘΥΝΣΗ]', general_directorate)
        
        paragraph.text = text
    
    # Αντικατάστασης στους πίνακες
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                
                if 'ΗΜΕΡΟΜΗΝΙΑ' in text:
                    text = text.replace('ΗΜΕΡΟΜΗΝΙΑ', report_date_gr)
                if '[ΓΕΝΙΚΗ_ΔΙΕΥΘΥΝΣΗ]' in text:
                    text = text.replace('[ΓΕΝΙΚΗ_ΔΙΕΥΘΥΝΣΗ]', general_directorate)
                
                cell.text = text
    
    # Αποθήκευση
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ DOCX αναφορά: {output_path}")


def create_xlsx_report(template_path, output_path, categorized_data, monday_str, sunday_str, charges_by_pkm=None):
    """
    Δημιουργεί XLSX από το πρότυπο με αναλυτικά δεδομένα
    
    Args:
        charges_by_pkm: Dict mapping PKM → charge record (for employee assignment)
    """
    import re
    
    if charges_by_pkm is None:
        charges_by_pkm = {}
    # Φόρτωση προτύπου
    wb = load_workbook(template_path)
    
    # Format week dates for title (DD-MM-YYYY)
    monday_date = datetime.strptime(monday_str, "%Y-%m-%d")
    sunday_date = datetime.strptime(sunday_str, "%Y-%m-%d")
    monday_formatted = monday_date.strftime("%d-%m-%Y")
    sunday_formatted = sunday_date.strftime("%d-%m-%Y")
    title_suffix = f"(ΑΠΟ {monday_formatted} ΕΩΣ {sunday_formatted})"
    
    # Καθαρισμός υπάρχοντος δεδομένων και προσθήκη νέων
    if 'ΝΕΑ' in wb.sheetnames:
        ws_new = wb['ΝΕΑ']
        
        # Ενημέρωση τίτλου (σειρά 1) με ημερομηνίες εβδομάδας
        title_cell = ws_new['A1']
        if title_cell.value:
            # Replace existing date range or append if not found
            original_title = str(title_cell.value)
            # Remove old date range if exists (pattern: (ΑΠΟ ... ΕΩΣ ...))
            base_title = re.sub(r'\(ΑΠΟ.*?ΕΩΣ.*?\)', '', original_title).strip()
            title_cell.value = f"{base_title}\n{title_suffix}"
        
        # Διατήρηση κεφαλίδας (πρώτες 2 σειρές)
        for row in ws_new.iter_rows(min_row=3, max_row=ws_new.max_row):
            for cell in row:
                cell.value = None
        
        # Προσθήκη νέων αιτήσεων
        row_idx = 3
        for idx, rec in enumerate(categorized_data['new_assigned'] + categorized_data['new_unassigned'], 1):
            ws_new[f'A{row_idx}'] = idx
            ws_new[f'B{row_idx}'] = rec.get('case_id', '')
            
            # Format ημερομηνίας μέχρι δευτερόλεπτα
            submitted_at = rec.get('submitted_at', '')
            if submitted_at and '.' in submitted_at:
                # Αφαίρεση microseconds: "2026-02-03 18:13:35.727528" -> "2026-02-03 18:13:35"
                submitted_at = submitted_at.split('.')[0]
            ws_new[f'C{row_idx}'] = submitted_at
            
            ws_new[f'D{row_idx}'] = rec.get('document_category', '')
            ws_new[f'E{row_idx}'] = rec.get('party', '')
            ws_new[f'F{row_idx}'] = rec.get('subject', '')
            
            # Ανάθεση σε: Βάση χρεώσεων (employee name ή "Χωρίς Ανάθεση")
            charge_info = rec.get('_charge', {})
            employee = charge_info.get('employee', '') if charge_info.get('charged') else 'Χωρίς Ανάθεση'
            ws_new[f'G{row_idx}'] = employee
            
            ws_new[f'H{row_idx}'] = rec.get('directory', '')
            ws_new[f'I{row_idx}'] = rec.get('department', '')
            # Αρ. Πρωτ Υπηρ -> αριθμός υπόθεσης (case_id)
            ws_new[f'J{row_idx}'] = rec.get('case_id', '')
            # Αρ. Πρωτ. (ΠΚΜ) -> αριθμός πρωτοκόλλου + έτος + Δ/νση
            protocol_number = rec.get('protocol_number', '')
            submitted_at_for_year = rec.get('submitted_at', '')
            year_part = ''
            if submitted_at_for_year:
                year_part = submitted_at_for_year[:4]
            directory = rec.get('directory', '')
            if protocol_number:
                ar_prot = f"{protocol_number} / {year_part} / {directory}" if year_part or directory else str(protocol_number)
            else:
                ar_prot = ''
            ws_new[f'K{row_idx}'] = ar_prot
            ws_new[f'L{row_idx}'] = rec.get('protocol_date', '')  # Ημ/νία Πρωτ
            
            row_idx += 1
    
    if 'ΠΑΛΙΑ' in wb.sheetnames:
        ws_old = wb['ΠΑΛΙΑ']
        
        # Ενημέρωση τίτλου (σειρά 1) με ημερομηνίες εβδομάδας
        title_cell = ws_old['A1']
        if title_cell.value:
            # Replace existing date range or append if not found
            original_title = str(title_cell.value)
            # Remove old date range if exists (pattern: (ΑΠΟ ... ΕΩΣ ...))
            base_title = re.sub(r'\(ΑΠΟ.*?ΕΩΣ.*?\)', '', original_title).strip()
            title_cell.value = f"{base_title}\n{title_suffix}"
        
        # Διατήρηση κεφαλίδας (πρώτες 2 σειρές)
        for row in ws_old.iter_rows(min_row=3, max_row=ws_old.max_row):
            for cell in row:
                cell.value = None
        
        # Προσθήκη παλιών αιτήσεων
        row_idx = 3
        for idx, rec in enumerate(categorized_data['old'], 1):
            ws_old[f'A{row_idx}'] = idx
            ws_old[f'B{row_idx}'] = rec.get('case_id', '')
            
            # Format ημερομηνίας μέχρι δευτερόλεπτα
            submitted_at = rec.get('submitted_at', '')
            if submitted_at and '.' in submitted_at:
                # Αφαίρεση microseconds: "2026-02-03 18:13:35.727528" -> "2026-02-03 18:13:35"
                submitted_at = submitted_at.split('.')[0]
            ws_old[f'C{row_idx}'] = submitted_at
            
            ws_old[f'D{row_idx}'] = rec.get('document_category', '')
            ws_old[f'E{row_idx}'] = rec.get('party', '')
            ws_old[f'F{row_idx}'] = rec.get('subject', '')
            
            # Ανάθεση σε: Βάση χρεώσεων (employee name ή "Χωρίς Ανάθεση")
            charge_info = rec.get('_charge', {})
            employee = charge_info.get('employee', '') if charge_info.get('charged') else 'Χωρίς Ανάθεση'
            ws_old[f'G{row_idx}'] = employee
            
            ws_old[f'H{row_idx}'] = rec.get('directory', '')
            ws_old[f'I{row_idx}'] = rec.get('department', '')
            # Αρ. Πρωτ Υπηρ -> αριθμός υπόθεσης (case_id)
            ws_old[f'J{row_idx}'] = rec.get('case_id', '')
            # Αρ. Πρωτ. (ΠΚΜ) -> αριθμός πρωτοκόλλου + έτος + Δ/νση
            protocol_number = rec.get('protocol_number', '')
            submitted_at_for_year = rec.get('submitted_at', '')
            year_part = ''
            if submitted_at_for_year:
                year_part = submitted_at_for_year[:4]
            directory = rec.get('directory', '')
            if protocol_number:
                ar_prot = f"{protocol_number} / {year_part} / {directory}" if year_part or directory else str(protocol_number)
            else:
                ar_prot = ''
            ws_old[f'K{row_idx}'] = ar_prot
            ws_old[f'L{row_idx}'] = rec.get('protocol_date', '')  # Ημ/νία Πρωτ
            
            row_idx += 1
    
    # Αποθήκευση
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    wb.save(output_path)
    print(f"✅ XLSX αναφορά: {output_path}")


def main():
    parser = argparse.ArgumentParser(
        description='Δημιουργία εβδομαδιαίας αναφοράς για ΣΗΔΕ αιτήσεις'
    )
    parser.add_argument(
        '--date',
        required=True,
        help='Ημερομηνία αναφοράς (YYYY-MM-DD). Παίρνει την προηγούμενη εβδομάδα.'
    )
    
    args = parser.parse_args()
    
    try:
        # Κανονικοποίηση ημερομηνίας: 2026-2-9 -> 2026-02-09
        date_str = args.date
        
        # Δοκιμή parsing με διάφορες μορφές
        parsed_date = None
        for fmt in ["%Y-%m-%d", "%Y-%-m-%-d"]:  # Windows δεν υποστηρίζει %-d
            try:
                parsed_date = datetime.strptime(date_str, fmt)
                break
            except ValueError:
                pass
        
        # Αν δεν ταιριάζει με τις παραπάνω μορφές, δοκίμασε manual parsing
        if not parsed_date:
            parts = date_str.split('-')
            if len(parts) == 3:
                try:
                    parsed_date = datetime(int(parts[0]), int(parts[1]), int(parts[2]))
                except (ValueError, IndexError):
                    pass
        
        if not parsed_date:
            print(f"❌ Λάθος μορφή ημερομηνίας: {args.date}. Χρησιμοποίησε YYYY-MM-DD (π.χ. 2026-02-09 ή 2026-2-9)")
            sys.exit(1)
        
        # Αν η ημερομηνία δεν είναι Δευτέρα, βρες την προηγούμενη Δευτέρα
        # Monday = 0, Tuesday = 1, ..., Sunday = 6
        day_of_week = parsed_date.weekday()
        if day_of_week != 0:  # Δεν είναι Δευτέρα
            days_to_subtract = day_of_week
            parsed_date = parsed_date - timedelta(days=days_to_subtract)
            print(f"ℹ️  Η δοθείσα ημερομηνία ({args.date}) δεν είναι Δευτέρα.")
            print(f"    Μετατόπιση στην προηγούμενη Δευτέρα: {parsed_date.strftime('%Y-%m-%d')}")
        
        # Κανονικοποίηση σε YYYY-MM-DD
        report_date = parsed_date.strftime("%Y-%m-%d")
    except Exception as e:
        print(f"❌ Σφάλμα επεξεργασίας ημερομηνίας: {e}")
        sys.exit(1)
    
    # Υπολογισμός εβδομάδας
    monday, sunday = get_week_boundaries(report_date)
    print(f"\n📋 Δημιουργία αναφοράς για εβδομάδα {monday} έως {sunday}")
    
    # Φόρτωση δεδομένων από API
    project_root = get_project_root()
    config = load_config(os.path.join(project_root, 'config', 'config.yaml'))
    monitor = PKMMonitor(
        base_url=config.get('base_url', 'https://shde.pkm.gov.gr'),
        urls=config.get('urls', {}),
        api_params=config.get('api_params', {}),
        login_params=config.get('login_params', {}),
        check_interval=config.get('check_interval', 300),
        username=config.get('username'),
        password=config.get('password'),
        session_cookies=config.get('session_cookies'))

    _ensure_logged_in(monitor)
    records = fetch_incoming_from_api(monitor, config)
    if not records:
        print("⚠️  Δεν βρέθηκαν εισερχόμενες αιτήσεις από το API.")
        sys.exit(1)
    
    # Φιλτράρισμα: Εξαίρεση αιτήσεων που υποβλήθηκαν ΜΕΤΑ το τέλος της εβδομάδας
    # Αυτό είναι κρίσιμο όταν παράγουμε αναφορά για παλιές εβδομάδες
    records = filter_records_by_week_end(records, sunday)
    if not records:
        print("⚠️  Δεν υπάρχουν αιτήσεις για την επιλεγμένη εβδομάδα.")
        sys.exit(1)
    
    # Προσπάθεια enrichment για όλα τα records ώστε να συμπληρωθούν τα πεδία λεπτομερειών
    # Αυτό θα προσπαθήσει να παρουσιάσει τα εξής πεδία μέσω detail API:
    # - protocol_number, protocol_date, procedure, directory, general_directorate, department
    try:
        to_enrich_all = [r for r in records if r.get('doc_id')]
        if to_enrich_all:
            print(f"🔄 Προσπάθεια enrichment {len(to_enrich_all)} εγγραφών...")
            enrich_record_details(monitor, to_enrich_all)
            print(f"✅ Enrichment ολοκληρώθηκε")
    except Exception as e:
        print(f"⚠️  Αποτυχία enrichment (συνεχίζουμε με τα διαθέσιμα δεδομένα): {e}")
    # Ανάκτηση χρεώσεων και εμπλουτισμός records
    charges_by_pkm = {}
    try:
        from charges import fetch_charges_combined, add_charge_info_from_combined
        print(f"📋 Ανάκτηση χρεώσεων υπαλλήλων (queryId=2 + queryId=3)...")
        charges_records, charges_by_pkm = fetch_charges_combined(monitor)
        print(f"   Βρέθηκαν {len(charges_records)} χρεώσεις από συνδυασμένες πηγές")
        records = add_charge_info_from_combined(records, charges_by_pkm, monitor=monitor, enrich_missing=True)
        print(f"   ✅ Εμπλουτισμός με χρεώσεις ολοκληρώθηκε (με API enrichment)")
    except Exception as exc:
        print(f"   ⚠️  Αποτυχία ανάκτησης χρεώσεων: {exc}")
        # Συνεχίζουμε χωρίς χρεώσεις
    groups = group_by_general_directorate(records)
    print(f"✅ Φορτώθηκαν {len(records)} εγγραφές. Γενικές Δ/νσεις: {len(groups)}")
    
    # Φιλτράρισμα διεκπεραιωμένων: Αφαίρεση αιτήσεων που ήδη διεκπεραιώθηκαν ΠΡΙΝ την ημερομηνία αναφοράς
    print(f"\n🔄 Φιλτράρισμα διεκπεραιωμένων αιτήσεων (ημερομηνία αναφοράς={report_date})...")
    filtered_groups = {}
    total_removed = 0
    total_settled_this_week = 0
    
    for general_directorate, recs in groups.items():
        filtered_recs = filter_out_settled_from_records(recs, monday, general_directorate, report_date)
        
        # Μέτρηση αφαιρεθεισών
        removed = len(recs) - len(filtered_recs)
        total_removed += removed
        
        # Μέτρηση νέων που διεκπεραιώθηκαν αυτή την εβδομάδα
        settled_this_week = len([r for r in filtered_recs if r.get('_settled_status') == 'Διεκπεραιωμένη'])
        total_settled_this_week += settled_this_week
        
        if filtered_recs:
            filtered_groups[general_directorate] = filtered_recs
    
    if total_removed > 0:
        print(f"   ✅ Αφαιρέθησαν {total_removed} διεκπεραιωμένες αιτήσεις (διεκπ. ΠΡΙΝ την {report_date})")
    if total_settled_this_week > 0:
        print(f"   ✅ {total_settled_this_week} νέες αιτήσεις διεκπεραιώθηκαν στη ίδια εβδομάδα (θα αφαιρεθούν από ΝΕΑ)")
    
    groups = filtered_groups
    if not groups:
        print("⚠️  Δεν υπάρχουν αιτήσεις για την επιλεγμένη περίοδο μετά το φιλτράρισμα διεκπεραιωμένων.")
        sys.exit(0)

    # Paths
    templates_dir = os.path.join(project_root, 'data', 'templates')
    outputs_dir = os.path.join(project_root, 'data', 'outputs', 'reports')

    # Εύρεση προτύπων
    docx_template = None
    xlsx_template = None

    for file in os.listdir(templates_dir):
        if file.endswith('.docx'):
            docx_template = os.path.join(templates_dir, file)
        elif file.endswith('.xlsx'):
            xlsx_template = os.path.join(templates_dir, file)

    if not docx_template or not xlsx_template:
        print(f"❌ Δεν βρέθηκαν πρότυπα στο {templates_dir}")
        sys.exit(1)

    # Εμφάνιση πίνακα στατιστικών
    print_statistics_table(groups, monday, sunday)

    # Δημιουργία αναφορών ανά Γενική Δ/νση
    print(f"🔄 Δημιουργία αναφορών ανά Γενική Δ/νση...")
    for general_directorate, recs in sorted(groups.items()):
        categorized = categorize_records(recs, monday, sunday)
        print(f"  • {general_directorate}: {len(categorized['all'])} εγγραφές")
        
        name_part = sanitize_filename_component(general_directorate)
        docx_filename = f"{report_date}_{name_part}_Αναφορά Εισερχομένων Αιτήσεων ΣΗΔΕ.docx"
        xlsx_filename = f"{report_date}_{name_part}_Αναφορά Εισερχομένων Αιτήσεων ΣΗΔΕ.xlsx"

        docx_path = os.path.join(outputs_dir, docx_filename)
        xlsx_path = os.path.join(outputs_dir, xlsx_filename)

        create_docx_report(docx_template, docx_path, report_date, categorized, general_directorate)
        create_xlsx_report(xlsx_template, xlsx_path, categorized, monday, sunday, charges_by_pkm)

    print(f"\n✨ Αναφορές δημιουργήθηκαν με επιτυχία!")


if __name__ == '__main__':
    main()
